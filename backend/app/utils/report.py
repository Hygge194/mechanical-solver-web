from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io

def format_num(val, decimals=2):
    try:
        return f"{float(val):.{decimals}f}"
    except (ValueError, TypeError):
        return str(val) if val is not None else "-"

def generate_report(results: dict):
    """
    Hàm để xuất file Word thuyết minh tự động và trả về file binary stream (BytesIO).
    `results` là JSON payload PROJECT_DATA từ Frontend.
    """
    doc = Document()
    
    # Tiêu đề lớn
    title = doc.add_heading('HỒ SƠ TÍNH TOÁN THIẾT KẾ HỆ DẪN ĐỘNG THÙNG TRỘN', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    
    # -- PHẦN 1: Thông tin chung & Dữ liệu đầu vào --
    doc.add_heading('Phần 1: Thông tin chung & Dữ liệu đầu vào', level=1)
    m1_data = results.get("M1_Data", {})
    input_data = m1_data.get("input", {})
    
    doc.add_paragraph(f"Công suất yêu cầu P: {format_num(input_data.get('pt'))} kW")
    doc.add_paragraph(f"Số vòng quay yêu cầu n: {format_num(input_data.get('niv'))} vòng/phút")
    doc.add_paragraph(f"Tuổi thọ làm việc L: {format_num(input_data.get('lifetime'))} giờ")
    doc.add_paragraph("Hãng ổ lăn lựa chọn: SKF (mặc định)")
    doc.add_paragraph()

    # -- PHẦN 2: Kết quả tính toán Động lực học (Module M1) --
    doc.add_heading('Phần 2: Kết quả tính toán Động lực học (Module M1)', level=1)
    if m1_data:
        motor = m1_data.get("motor", {})
        kin = m1_data.get("kinematics", {})
        ratios = m1_data.get("ratios", {})
        
        doc.add_heading('Thông số động cơ được chọn:', level=2)
        doc.add_paragraph(f"Mã động cơ: {motor.get('code', motor.get('model', '-'))}")
        doc.add_paragraph(f"Công suất định mức (Pđm): {format_num(motor.get('P'))} kW")
        doc.add_paragraph(f"Tốc độ quay định mức (nđm): {format_num(motor.get('n', '-'), 0)} vòng/phút")
        eta = motor.get('eta', motor.get('HieuSuat', '-'))
        doc.add_paragraph(f"Hiệu suất (η): {format_num(eta, 3)}")
        
        doc.add_heading('Tỷ số truyền:', level=2)
        doc.add_paragraph(f"Tỷ số truyền tổng u_tổng: {format_num(ratios.get('u_t'), 3)}")
        doc.add_paragraph(f"Tỷ số truyền đai u_đai: {format_num(ratios.get('u_dai'), 3)}")
        doc.add_paragraph(f"Tỷ số truyền bánh răng côn u_bánh_răng_côn: {format_num(ratios.get('u_1'), 3)}")
        doc.add_paragraph(f"Tỷ số truyền bánh răng trụ u_bánh_răng_trụ: {format_num(ratios.get('u_2'), 3)}")
        
        doc.add_heading('Bảng thông số động lực học các trục:', level=2)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Trục'
        hdr_cells[1].text = 'Công suất P (kW)'
        hdr_cells[2].text = 'Vận tốc quay n (vòng/phút)'
        hdr_cells[3].text = 'Mô-men xoắn T (N.mm)'
        
        for key, name in [('truc_dc', 'Trục động cơ'), ('truc_1', 'Trục I'), ('truc_2', 'Trục II'), ('truc_3', 'Trục III/Trục công tác')]:
            if key in kin and kin[key]:
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = format_num(kin[key].get('P'), 3)
                row_cells[2].text = format_num(kin[key].get('n'), 1)
                row_cells[3].text = format_num(kin[key].get('T'), 3)
    else:
        doc.add_paragraph("Không có dữ liệu Module 1.")
    doc.add_paragraph()

    # -- PHẦN 3: Chi tiết thiết kế các bộ truyền (Module M2, M3, M4) --
    doc.add_heading('Phần 3: Chi tiết thiết kế các bộ truyền (Module M2, M3, M4)', level=1)
    
    # M2
    m2_data = results.get("M2_Data", {}).get("result", {})
    if not m2_data and "belt_type" in results.get("M2_Data", {}):
        m2_data = results.get("M2_Data", {})
        
    doc.add_heading('Bộ truyền đai (M2):', level=2)
    if m2_data:
        doc.add_paragraph(f"Loại đai hình thang: {m2_data.get('belt_type', '-')}")
        doc.add_paragraph(f"Đường kính bánh đai d1, d2: {format_num(m2_data.get('d1'))} mm, {format_num(m2_data.get('d2'))} mm")
        doc.add_paragraph(f"Chiều dài đai L: {format_num(m2_data.get('L'))} mm")
        doc.add_paragraph(f"Khoảng cách trục a: {format_num(m2_data.get('a'))} mm")
        doc.add_paragraph(f"Số lượng dây đai Z: {format_num(m2_data.get('Z'), 0)}")
        doc.add_paragraph(f"Lực căng đai ban đầu F0: {format_num(m2_data.get('F0'))} N")
        doc.add_paragraph(f"Lực tác dụng lên trục Fr: {format_num(m2_data.get('Fr'))} N")
    else:
        doc.add_paragraph("Không có dữ liệu Module 2.")
        
    # M3
    m3_data = results.get("M3_Data", {})
    doc.add_heading('Bộ truyền bánh răng côn (M3):', level=2)
    if m3_data:
        doc.add_paragraph(f"Vật liệu: {m3_data.get('material', 'Thép hợp kim')}")
        doc.add_paragraph(f"Chiều dài côn ngoài Re: {format_num(m3_data.get('Re'))} mm")
        doc.add_paragraph(f"Mô-đun m: {format_num(m3_data.get('mte'))} mm")
        doc.add_paragraph(f"Số răng z1, z2: {format_num(m3_data.get('z1'), 0)}, {format_num(m3_data.get('z2'), 0)}")
        doc.add_paragraph(f"Góc côn δ1, δ2: {format_num(m3_data.get('delta1_deg'))}°, {format_num(m3_data.get('delta2_deg'))}°")
        doc.add_paragraph(f"Đường kính vòng chia de1, de2: {format_num(m3_data.get('de1'))} mm, {format_num(m3_data.get('de2'))} mm")
        doc.add_paragraph(f"Ứng suất tiếp xúc σH: {format_num(m3_data.get('sigH'))} MPa")
        doc.add_paragraph(f"Ứng suất uốn σF1, σF2: {format_num(m3_data.get('sigF1'))} MPa, {format_num(m3_data.get('sigF2'))} MPa")
        doc.add_paragraph(f"Kết luận kiểm nghiệm: {m3_data.get('status', '-')}")
    else:
        doc.add_paragraph("Không có dữ liệu Module 3.")
        
    # M4
    m4_data = results.get("M4_Data", {})
    doc.add_heading('Bộ truyền bánh răng trụ (M4):', level=2)
    if m4_data:
        doc.add_paragraph("Loại răng: Răng thẳng")
        doc.add_paragraph("Góc nghiêng β: 0°")
        doc.add_paragraph(f"Khoảng cách trục aw: {format_num(m4_data.get('aw'))} mm")
        doc.add_paragraph(f"Mô-đun m: {format_num(m4_data.get('m'))} mm")
        doc.add_paragraph(f"Số răng z1, z2: {format_num(m4_data.get('z1'), 0)}, {format_num(m4_data.get('z2'), 0)}")
        doc.add_paragraph(f"Ứng suất tiếp xúc σH: {format_num(m4_data.get('sigmaH'))} MPa")
        doc.add_paragraph(f"Ứng suất uốn σF1, σF2: {format_num(m4_data.get('sigmaF1'))} MPa, {format_num(m4_data.get('sigmaF2'))} MPa")
        doc.add_paragraph(f"Kết luận kiểm nghiệm: {m4_data.get('status', '-')}")
        
        doc.add_heading('Lực tác dụng lên bộ truyền bánh răng trụ:', level=3)
        doc.add_paragraph(f"Lực vòng Ft: {format_num(m4_data.get('Ft'))} N")
        doc.add_paragraph(f"Lực hướng tâm Fr: {format_num(m4_data.get('Fr'))} N")
        doc.add_paragraph(f"Lực dọc trục Fa: {format_num(m4_data.get('Fa'))} N")
    else:
        doc.add_paragraph("Không có dữ liệu Module 4.")

    doc.add_paragraph()
    doc.add_paragraph("---- Hết báo cáo ----")
    
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream
